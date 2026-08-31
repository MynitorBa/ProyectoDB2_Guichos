"""Genera los tres documentos de documentación del proyecto TiendaYa."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_heading_color(paragraph, r, g, b):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(r, g, b)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    set_heading_color(p, 31, 78, 121)
    doc.add_paragraph()


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    set_heading_color(p, 17, 122, 101)


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    set_heading_color(p, 88, 88, 88)


def body(doc, text):
    doc.add_paragraph(text)


def bullet(doc, text, level=0):
    doc.add_paragraph(text, style='List Bullet')


def separador(doc):
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO 1 — DOCUMENTACIÓN DEL CÓDIGO
# ══════════════════════════════════════════════════════════════════════════════
def crear_doc1():
    doc = Document()
    doc.core_properties.title = "Documentación del Código - TiendaYa"

    # Portada
    titulo = doc.add_heading("TiendaYa — Documentación del Código", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Descripción detallada de cada archivo del frontend y el backend")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ──────────────────────────────────────────────────
    h1(doc, "PARTE 1 — FRONTEND (React + Vite)")
    body(doc, "El frontend es una aplicación web construida con React y Vite. Se organiza en páginas, componentes reutilizables, módulos de comunicación con el servidor (API) y contextos de estado global. Todo el código fuente vive en la carpeta frontend/src/.")
    separador(doc)

    # ── Punto de entrada ──
    h2(doc, "Archivos de Arranque")

    h3(doc, "main.jsx — Punto de entrada de la aplicación")
    body(doc, "Es el primer archivo que se ejecuta cuando se abre la aplicación en el navegador. Su trabajo es montar toda la estructura de proveedores de estado antes de mostrar cualquier pantalla. Envuelve la aplicación con QueryClientProvider (que gestiona las peticiones al servidor y su caché), luego con AuthProvider (que sabe si hay un usuario conectado), luego con CartProvider (que mantiene el estado del carrito de compras), y finalmente renderiza el componente App. También configura el sistema de notificaciones emergentes (toasts) que aparecen en la esquina de la pantalla.")
    separador(doc)

    h3(doc, "App.jsx — Estructura visual principal")
    body(doc, "Define la estructura visual permanente de cada pantalla: un encabezado fijo arriba, el contenido dinámico en el centro (que cambia según la página) y un pie de página abajo. No contiene lógica de negocio; solo organiza el espacio visual.")
    separador(doc)

    h3(doc, "router.jsx — Definición de rutas")
    body(doc, "Declara qué componente se muestra para cada dirección URL. Tiene dos tipos de rutas: las públicas (cualquier visitante puede verlas, como el catálogo o el inicio de sesión) y las privadas (requieren estar conectado). Las rutas privadas también verifican el rol del usuario: solo los administradores pueden entrar al panel de administración y solo los vendedores al panel de vendedor. Si alguien intenta acceder sin permiso, se le redirige automáticamente.")
    bullet(doc, "Rutas públicas: inicio, catálogo, detalle de producto, login, registro")
    bullet(doc, "Rutas privadas de comprador: carrito, checkout, pedidos, perfil")
    bullet(doc, "Rutas privadas de vendedor: panel de vendedor")
    bullet(doc, "Rutas privadas de administrador: panel de administración")
    separador(doc)

    h3(doc, "index.css — Estilos globales")
    body(doc, "Define la paleta de colores del sitio usando variables CSS (por ejemplo --color-action para el color azul principal). También establece tipografías: Plus Jakarta Sans para títulos, Inter para texto normal y JetBrains Mono para números y códigos. Incluye soporte para modo oscuro completo: si el usuario activa el modo oscuro en su sistema operativo, todos los colores cambian automáticamente.")
    separador(doc)

    # ── API ──
    h2(doc, "Módulos de Comunicación con el Servidor (src/api/)")
    body(doc, "Cada archivo en esta carpeta representa un grupo de funciones que se comunican con el servidor. Todas las funciones envían peticiones HTTP y devuelven los datos que el servidor responde.")
    separador(doc)

    h3(doc, "client.js — Cliente HTTP base")
    body(doc, "Crea una instancia de Axios (la librería que se usa para hacer peticiones HTTP) preconfigurada con la dirección base del servidor. Hace dos cosas importantes de forma automática. Primero, antes de cada petición, revisa si hay un token de sesión guardado en el navegador y lo agrega al encabezado de la petición para que el servidor sepa quién está haciendo la solicitud. Segundo, si el servidor responde con un error 401 (sesión vencida o inválida), automáticamente borra el token guardado y redirige al usuario a la pantalla de inicio de sesión. Todos los demás módulos de API usan este cliente como base.")
    separador(doc)

    h3(doc, "auth.js — Autenticación")
    body(doc, "Contiene las funciones relacionadas con el acceso al sistema. La función login envía el correo y la contraseña al servidor y recibe un token de acceso. La función register crea una nueva cuenta. La función me consulta los datos del usuario actualmente conectado. La función updateProfile actualiza el nombre, apellido y teléfono del perfil.")
    separador(doc)

    h3(doc, "products.js — Productos y catálogo")
    body(doc, "Agrupa todas las funciones relacionadas con productos. Las funciones públicas (que no requieren login) incluyen getProducts para obtener el listado del catálogo con filtros de categoría, precio y búsqueda por texto, y getCategories para listar las categorías disponibles. Las funciones administrativas permiten crear, actualizar y eliminar productos, consultar el historial de cambios de un producto y ver el estado que tenía en una fecha específica.")
    separador(doc)

    h3(doc, "cart.js — Carrito de compras")
    body(doc, "Tres funciones simples: getCart obtiene el contenido actual del carrito del servidor, addItem agrega una oferta al carrito indicando la cantidad, y removeItem elimina un artículo del carrito. El carrito vive en el servidor (no solo en el navegador) para que persista si el usuario cierra y vuelve a abrir el sitio.")
    separador(doc)

    h3(doc, "orders.js — Pedidos y checkout")
    body(doc, "Maneja todo el proceso de compra. La función checkout envía la dirección de entrega, el método de pago y los artículos al servidor para crear el pedido. Las funciones getOrders y getOrder consultan el historial de pedidos. La función getOrderInvoice descarga la factura en formato PDF. También incluye funciones para gestionar las direcciones de entrega guardadas (agregar, actualizar, eliminar).")
    separador(doc)

    h3(doc, "notifications.js — Notificaciones")
    body(doc, "Permite obtener las notificaciones del usuario (getNotifications), contar cuántas no han sido leídas (getUnreadCount), y marcarlas como leídas individualmente o todas a la vez.")
    separador(doc)

    h3(doc, "vendor.js — Panel de vendedor")
    body(doc, "Funciones exclusivas para usuarios con rol de vendedor. Incluye getVendorStats para ver estadísticas del negocio, getVendorOrders para ver los pedidos que contienen productos del vendedor, updateVendorOrderStatus para cambiar el estado de un subpedido, y funciones para proponer nuevos productos o solicitar ser el oferente de un producto existente.")
    separador(doc)

    h3(doc, "admin.js — Panel de administración")
    body(doc, "El módulo más grande. Contiene docenas de funciones exclusivas para administradores que cubren la gestión completa de usuarios, productos, categorías, pedidos, ventas y solicitudes de catálogo. Permite exportar reportes de ventas a Excel, cambiar roles de usuarios, aprobar o rechazar propuestas de vendedores, y consultar estadísticas globales del sistema.")
    separador(doc)

    # ── Contextos ──
    h2(doc, "Estado Global (src/context/)")
    body(doc, "Los contextos de React permiten que cualquier componente de la aplicación acceda a información compartida sin tener que pasarla manualmente de componente en componente.")
    separador(doc)

    h3(doc, "AuthContext.jsx — Estado de autenticación")
    body(doc, "Guarda si hay un usuario conectado y cuáles son sus datos (nombre, correo, roles). Al cargar la aplicación, revisa si hay un token guardado en el navegador y, si lo hay, lo valida consultando al servidor para obtener los datos actualizados del usuario. Provee tres funciones: signIn para guardar el token y los datos cuando el usuario inicia sesión, signOut para cerrar la sesión y limpiar todos los datos guardados, y updateUser para actualizar los datos del perfil sin cerrar sesión.")
    separador(doc)

    h3(doc, "CartContext.jsx — Estado del carrito")
    body(doc, "Mantiene el contenido del carrito disponible en toda la aplicación. Las funciones fetchCart, add y remove sincronizan el carrito local con el servidor cada vez que hay un cambio. Si el usuario no está conectado, las operaciones del carrito simplemente no hacen nada de forma silenciosa.")
    separador(doc)

    # ── Utilidades ──
    h2(doc, "Utilidades (src/lib/)")

    h3(doc, "utils.js — Funciones de apoyo")
    body(doc, "Contiene funciones pequeñas que se usan en toda la aplicación. La función cn combina clases de estilos de forma inteligente, resolviendo conflictos entre ellas. La función formatQ convierte un número a formato de quetzales guatemaltecos con dos decimales. La función formatDate muestra fechas en formato día/mes/año en español. La función slugify convierte un texto a formato de URL (minúsculas, sin espacios).")
    separador(doc)

    h3(doc, "queryClient.js — Configuración de caché")
    body(doc, "Crea el objeto de configuración para React Query, la librería que gestiona las peticiones al servidor. Configura que los datos se consideren 'frescos' durante 30 segundos (evitando peticiones innecesarias) y que solo se reintente una vez si una petición falla.")
    separador(doc)

    # ── Componentes UI ──
    h2(doc, "Componentes de Interfaz Reutilizables (src/components/ui/)")
    body(doc, "Son los bloques de construcción visuales básicos de la aplicación. Cada uno es un componente pequeño y reutilizable con un diseño consistente.")
    separador(doc)

    h3(doc, "button.jsx — Botón")
    body(doc, "El componente de botón con múltiples variantes de color (azul primario, secundario, rojo para acciones destructivas, transparente) y tamaños (pequeño, mediano, grande). Tiene soporte nativo para estado de carga, mostrando un spinner animado y desactivando el clic mientras se procesa una acción. También puede renderizarse como un enlace de navegación.")
    separador(doc)

    h3(doc, "badge.jsx — Etiqueta de estado")
    body(doc, "Pequeña etiqueta de colores usada para mostrar estados como 'activo', 'pendiente', 'cancelado'. Tiene variantes de color para cada tipo de estado: verde para éxito, amarillo para advertencias, rojo para errores.")
    separador(doc)

    h3(doc, "dialog.jsx — Ventana modal")
    body(doc, "Ventana emergente accesible que oscurece el fondo. Se usa para formularios de edición, confirmaciones y detalles. Incluye un botón de cierre automático y manejo de teclado (la tecla Escape la cierra).")
    separador(doc)

    h3(doc, "sheet.jsx — Panel lateral deslizante")
    body(doc, "Similar a un diálogo pero aparece deslizándose desde un lado de la pantalla (usualmente la derecha). Se usa para el carrito de compras rápido y los filtros en móvil.")
    separador(doc)

    h3(doc, "input.jsx — Campo de texto")
    body(doc, "Campo de texto y área de texto con estilos consistentes. Muestra un borde rojo y mensaje de error cuando la validación falla.")
    separador(doc)

    h3(doc, "select.jsx — Selector con búsqueda")
    body(doc, "Selector desplegable que incluye un campo de búsqueda para filtrar las opciones. Útil cuando hay muchas opciones (por ejemplo, para elegir una categoría).")
    separador(doc)

    h3(doc, "skeleton.jsx — Marcador de posición de carga")
    body(doc, "Muestra rectángulos animados mientras se cargan los datos reales. Evita que la pantalla se vea vacía durante la carga y da al usuario una indicación visual de la estructura que aparecerá.")
    separador(doc)

    h3(doc, "star-rating.jsx — Calificación con estrellas")
    body(doc, "Muestra la calificación de un producto con estrellas, soportando valores decimales (por ejemplo, 4.3 estrellas).")
    separador(doc)

    # ── Componentes específicos ──
    h2(doc, "Componentes de Funcionalidad Específica (src/components/)")
    separador(doc)

    h3(doc, "layout/Header.jsx — Encabezado global")
    body(doc, "El encabezado que aparece en todas las páginas. Contiene el logo de TiendaYa, una barra de búsqueda que redirige al catálogo con el término buscado, un selector de tema oscuro/claro, la campanita de notificaciones (que consulta el servidor cada 60 segundos para mostrar el número de notificaciones sin leer), el icono del carrito con el contador de artículos, el menú desplegable del usuario con acceso al perfil y al panel correspondiente según su rol, y un menú de categorías que muestra todas las categorías en un dropdown al hacer clic.")
    separador(doc)

    h3(doc, "cart/CartSheet.jsx — Carrito lateral rápido")
    body(doc, "Panel que se desliza desde la derecha al hacer clic en el icono del carrito. Muestra todos los artículos con imagen, nombre y precio, permite eliminar artículos directamente, muestra el desglose de subtotal, IVA incluido y total, y tiene botones para ir al checkout o a la página completa del carrito.")
    separador(doc)

    h3(doc, "product/ProductCard.jsx — Tarjeta de producto")
    body(doc, "La tarjeta visual que representa un producto en el catálogo. Muestra la imagen principal del producto, el nombre, el vendedor, la calificación con estrellas, el precio y un botón para agregar al carrito. Si el producto no tiene stock, el botón se deshabilita y aparece un badge de 'Sin stock'. Si hay pocas unidades (5 o menos), muestra una alerta de 'Solo X en stock'. Al agregar al carrito, el botón cambia temporalmente a 'Agregado' con un ícono de check para confirmar la acción.")
    separador(doc)

    h3(doc, "product/ProductImage.jsx — Imagen de producto con respaldo")
    body(doc, "Muestra la imagen de un producto y, si no carga por cualquier razón, muestra automáticamente un ícono representativo según la categoría del producto. Esto evita que aparezcan imágenes rotas en la interfaz.")
    separador(doc)

    h3(doc, "product/CategoryAttrPanel.jsx — Panel de atributos por categoría")
    body(doc, "Muestra los atributos específicos de cada tipo de producto de forma adaptada. Para electrónica muestra un grid con procesador, RAM y almacenamiento. Para ropa muestra selectores de talla y un visualizador de color. Para libros muestra autor, editorial e ISBN. Para alimentos muestra el peso, origen y etiquetas nutricionales. Para cualquier otro tipo muestra los atributos de forma genérica. Esto es posible porque la configuración de atributos viene de MongoDB y es diferente para cada categoría.")
    separador(doc)

    h3(doc, "admin/CatalogRequestsSection.jsx — Solicitudes de catálogo (admin)")
    body(doc, "Sección del panel de administración que muestra todas las solicitudes enviadas por vendedores para agregar productos o hacerse oferentes. Permite filtrar por estado (pendiente, aprobada, rechazada). Al hacer clic en 'Revisar', se abre una ventana con todos los detalles de la solicitud: producto propuesto, imágenes, precio, observaciones del vendedor. Si la solicitud está pendiente, el administrador puede aprobarla o rechazarla con un comentario.")
    separador(doc)

    h3(doc, "vendor/CatalogRequestsSection.jsx — Solicitudes de catálogo (vendedor)")
    body(doc, "Sección del panel del vendedor para gestionar sus propuestas. Tiene dos botones principales: uno para proponer un producto completamente nuevo (con nombre, descripción, categorías, atributos, imágenes y precio) y otro para solicitar vender un producto que ya existe en el catálogo (indicando el precio y stock propios). También muestra el historial de solicitudes enviadas con su estado actual.")
    separador(doc)

    # ── Páginas ──
    h2(doc, "Páginas (src/pages/)")
    body(doc, "Las páginas son los componentes que se muestran al navegar a una URL específica. Cada una orquesta componentes más pequeños para formar una pantalla completa.")
    separador(doc)

    h3(doc, "HomePage.jsx — Página de inicio")
    body(doc, "La primera pantalla que ven los visitantes. Tiene una sección hero con el lema y dos llamadas a la acción: explorar el catálogo y crear una cuenta. Debajo muestra un carrusel horizontal con todas las categorías que el usuario puede clicar para filtrar el catálogo. Luego un grid de productos destacados y finalmente una sección que explica las ventajas de comprar en TiendaYa.")
    separador(doc)

    h3(doc, "LoginPage.jsx — Inicio de sesión")
    body(doc, "Formulario simple con correo y contraseña. Al enviar, llama al servidor para validar las credenciales. Si son correctas, guarda el token de sesión y redirige al usuario a la página de inicio. Si son incorrectas, muestra un mensaje de error. Usa React Hook Form para validar que los campos no estén vacíos antes de enviar.")
    separador(doc)

    h3(doc, "RegisterPage.jsx — Registro de cuenta")
    body(doc, "Formulario de registro que en un solo paso recoge los datos personales (nombre, apellido, correo, contraseña) y la primera dirección de envío (departamento, municipio, dirección). La contraseña tiene validación en tiempo real con íconos que indican si cumple cada requisito: mínimo 8 caracteres, al menos una mayúscula, una minúscula, un número y un carácter especial. Al registrarse exitosamente, el sistema inicia la sesión automáticamente.")
    separador(doc)

    h3(doc, "CatalogPage.jsx — Catálogo de productos")
    body(doc, "El listado principal de productos con filtros. En la barra lateral (o en un panel deslizante en móvil) el usuario puede filtrar por categoría, rango de precio y ordenar los resultados. La búsqueda por texto viene del encabezado y se refleja en la URL para que se pueda compartir el enlace con los filtros aplicados. Los resultados se muestran en un grid de tarjetas de producto con paginación.")
    separador(doc)

    h3(doc, "ProductDetailPage.jsx — Detalle de producto")
    body(doc, "La página más compleja del frontend. Muestra una galería de imágenes con selector de miniaturas, el nombre y descripción del producto, el precio, la información del vendedor, un selector de variantes si el producto tiene varias (por ejemplo, diferentes capacidades de almacenamiento), un selector para elegir entre varios vendedores si hay más de uno ofreciendo el mismo producto, el selector de cantidad, el botón de agregar al carrito, los atributos técnicos específicos de la categoría y las reseñas de otros compradores con su distribución de estrellas.")
    separador(doc)

    h3(doc, "CartPage.jsx — Carrito de compras completo")
    body(doc, "Lista todos los artículos en el carrito con opción de eliminar cada uno. En la columna derecha muestra el resumen del pedido con la base imponible (precio sin IVA), el IVA del 12% extraído del precio (porque en Guatemala el IVA ya está incluido en el precio), y el total a pagar. Tiene un botón para proceder al proceso de pago.")
    separador(doc)

    h3(doc, "CheckoutPage.jsx — Proceso de pago")
    body(doc, "Guía al usuario por tres pasos para completar la compra. En el primer paso elige la dirección de entrega de entre sus direcciones guardadas. En el segundo paso elige el método de pago. En el tercer paso ve un resumen completo y confirma el pedido. Al confirmar, el servidor crea el pedido de forma transaccional y, si todo sale bien, muestra una pantalla de confirmación con el número de pedido. A la derecha en todos los pasos se muestra el resumen del pedido para no perder de vista el total.")
    separador(doc)

    h3(doc, "OrdersPage.jsx — Historial de pedidos")
    body(doc, "Muestra todos los pedidos realizados por el usuario, ordenados del más reciente al más antiguo. Cada pedido muestra su número, fecha, estado actual con un color indicativo y el total. Desde aquí se puede descargar la factura en PDF de cualquier pedido.")
    separador(doc)

    h3(doc, "OrderDetailPage.jsx — Detalle de pedido")
    body(doc, "Muestra el detalle completo de un pedido específico. Incluye una barra de progreso visual que indica en qué etapa está el envío (confirmado, enviado, entregado). Debajo muestra cada producto comprado con su precio unitario, cantidad y subtotal de línea. Al final muestra el desglose del pago con la base, el IVA y el total.")
    separador(doc)

    h3(doc, "ProfilePage.jsx — Perfil de usuario")
    body(doc, "Permite al usuario ver y editar sus datos personales (nombre, apellido, teléfono). También gestiona sus direcciones de envío guardadas: puede agregar nuevas, editar existentes y eliminar las que ya no necesita. El correo electrónico no se puede cambiar porque es el identificador único de la cuenta.")
    separador(doc)

    h3(doc, "VendorPage.jsx — Panel de vendedor")
    body(doc, "El espacio de trabajo para usuarios con rol de vendedor. Tiene tres secciones accesibles por pestañas: estadísticas de ventas (total de pedidos, ingresos y pedidos pendientes de preparar), gestión de pedidos (tabla con todos los pedidos que incluyen sus productos, con opción de cambiar el estado de cada uno) y solicitudes de catálogo (para proponer productos o unirse como oferente de productos existentes).")
    separador(doc)

    h3(doc, "AdminPage.jsx — Panel de administración")
    body(doc, "El panel central de gestión del sistema con seis secciones. Estadísticas muestra KPIs del catálogo y gráficas de precios. Productos permite el CRUD completo de productos con gestión de imágenes. Categorías permite definir las categorías y sus esquemas de atributos dinámicos. Usuarios permite ver y cambiar los roles de los usuarios del sistema. Pedidos permite ver y cambiar el estado de todos los pedidos. Ventas muestra reportes de ingresos con gráfica de tendencia y opción de exportar a Excel. Solicitudes muestra el flujo de aprobación de propuestas de vendedores.")
    separador(doc)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # BACKEND
    # ══════════════════════════════════════════════════
    h1(doc, "PARTE 2 — BACKEND (FastAPI + Python)")
    body(doc, "El backend es una API REST construida con FastAPI. Recibe peticiones del frontend, aplica las reglas de negocio y se comunica con MySQL y MongoDB para guardar y recuperar datos. Se organiza en cuatro capas: núcleo (configuración y conexiones), modelos (estructura de los datos en MySQL), esquemas (estructura de los datos en la API), servicios (lógica de negocio) y endpoints (la API pública).")
    separador(doc)

    h2(doc, "Núcleo del Sistema (app/core/)")
    separador(doc)

    h3(doc, "config.py — Configuración centralizada")
    body(doc, "Lee todas las variables de configuración del archivo .env (contraseñas, puertos, claves secretas). Si una variable no está definida, usa un valor por defecto seguro. Expone propiedades calculadas como mysql_url que construye la cadena de conexión completa a partir de sus partes. Toda la aplicación accede a la configuración a través del objeto settings importado de este archivo, así que cambiar una variable en .env afecta a todo el sistema.")
    separador(doc)

    h3(doc, "db_mysql.py — Conexión a MySQL")
    body(doc, "Crea y configura el motor de conexión a MySQL usando SQLAlchemy. Usa un pool de conexiones de hasta 10 conexiones activas simultáneas para soportar múltiples usuarios a la vez. La opción pool_pre_ping verifica que las conexiones sigan activas antes de usarlas, lo que evita errores cuando MySQL cierra conexiones inactivas. Provee la función get_db, un generador que entrega una sesión de base de datos a cada endpoint y garantiza que se cierre al terminar, aunque ocurra un error.")
    separador(doc)

    h3(doc, "db_mongo.py — Conexión a MongoDB")
    body(doc, "Mantiene una única conexión al servidor de MongoDB (patrón singleton) para que no se abran conexiones nuevas en cada petición. La función ensure_indexes crea todos los índices de las colecciones de MongoDB al arrancar la aplicación; como MongoDB crea índices de forma idempotente, llamar esta función varias veces no genera errores. Los índices cubren los patrones de búsqueda más frecuentes: filtrar productos por categoría y disponibilidad, buscar por texto en nombre y descripción, y reconstruir el historial de eventos de un producto.")
    separador(doc)

    h3(doc, "deps.py — Dependencias reutilizables")
    body(doc, "Define funciones que FastAPI inyecta automáticamente en los endpoints. La función get_current_user recibe el token JWT del encabezado de la petición, lo valida, y devuelve el usuario de MySQL correspondiente. Si el token no es válido o el usuario está inactivo, lanza un error 401. La función require_role genera dependencias de control de acceso: require_role('administrador') solo permite pasar a usuarios con ese rol, devolviendo un error 403 a los demás.")
    separador(doc)

    h3(doc, "security.py — Seguridad y tokens")
    body(doc, "Dos responsabilidades: gestión de contraseñas y gestión de tokens JWT. Para contraseñas usa bcrypt, que aplica una función de hash irreversible con sal aleatoria. Esto significa que aunque la base de datos sea comprometida, las contraseñas no pueden ser recuperadas. Para los tokens JWT, crea tokens firmados con una clave secreta que incluyen la identidad del usuario y una fecha de expiración. La función decode_token valida la firma y devuelve el contenido del token, o un diccionario vacío si es inválido.")
    separador(doc)

    h2(doc, "Modelos de Base de Datos MySQL (app/models/)")
    body(doc, "Cada archivo define una clase Python que representa una tabla de MySQL. SQLAlchemy usa estas clases para traducir operaciones Python a consultas SQL.")
    separador(doc)

    h3(doc, "usuario.py — Usuarios y roles")
    body(doc, "Define tres clases. Rol es la tabla de roles del sistema (comprador, vendedor, administrador). UsuarioRol es la tabla intermedia que une usuarios con roles (relación muchos-a-muchos). Usuario es la tabla principal con los datos del usuario: email, contraseña hasheada, nombre, apellido, teléfono, estado de la cuenta y si el email fue verificado. Un usuario puede tener múltiples roles.")
    separador(doc)

    h3(doc, "vendedor.py — Perfil comercial")
    body(doc, "Extensión opcional del usuario para quienes son vendedores. Tiene relación uno-a-uno con Usuario. Almacena el nombre comercial del negocio, el NIT (que debe ser único), una descripción, el logo, y el estado de verificación del perfil (pendiente, verificado o rechazado por el administrador).")
    separador(doc)

    h3(doc, "direccion.py — Direcciones de entrega")
    body(doc, "Cada fila es una dirección guardada por un usuario. Tiene tipo (envío o facturación), país, departamento, municipio y las líneas de dirección. La columna es_predeterminada indica cuál usar por defecto en el checkout. La columna activa permite el borrado lógico: la dirección no se elimina realmente de la base de datos, solo se marca como inactiva.")
    separador(doc)

    h3(doc, "categoria.py — Categorías jerárquicas")
    body(doc, "Árbol de categorías donde cada categoría puede tener una categoría padre. Las categorías sin padre son las raíces (por ejemplo, 'Electrónica'). Las categorías con padre son subcategorías (por ejemplo, 'Computadoras' dentro de 'Electrónica'). Tiene un slug único (texto en formato URL) y un prefijo para generar SKUs.")
    separador(doc)

    h3(doc, "producto_referencia.py — Ancla SQL del producto")
    body(doc, "Tabla mínima que representa la existencia de un producto en el sistema. Solo guarda el identificador del producto en MongoDB (producto_ref), la categoría y la fecha de creación. El detalle completo del producto (nombre, descripción, atributos, imágenes) vive en MongoDB. Esta tabla existe para poder referenciar productos desde otras tablas MySQL como ofertas e inventario.")
    separador(doc)

    h3(doc, "producto_variante_referencia.py — Ancla SQL de variante")
    body(doc, "Similar a la anterior pero para variantes dinámicas de un producto. Guarda el identificador de la variante en MongoDB (variante_ref) y a qué producto pertenece. El detalle de la variante (atributos específicos como capacidad de almacenamiento o color) vive en MongoDB en la colección producto_variantes.")
    separador(doc)

    h3(doc, "producto_imagen.py — Imágenes de productos")
    body(doc, "Almacena las imágenes de productos directamente en MySQL como datos binarios (LONGBLOB). Cada imagen tiene un orden que determina cuál aparece primero en la galería, quién la subió y cuándo.")
    separador(doc)

    h3(doc, "oferta.py — Ofertas de venta")
    body(doc, "Una oferta representa a un vendedor ofreciendo un producto específico a un precio. Un mismo producto puede tener múltiples ofertas de diferentes vendedores. Tiene el precio actual, el estado (borrador, activa, pausada, descontinuada) y un SKU único por vendedor. También incluye dos tablas de historial: OfertaPrecioHistorial registra todos los cambios de precio con sus fechas de vigencia (tipo SCD2), y OfertaEstadoHistorial hace lo mismo para los cambios de estado.")
    separador(doc)

    h3(doc, "inventario.py — Control de stock")
    body(doc, "El inventario se gestiona por oferta (no solo por producto), porque cada vendedor puede tener diferente stock del mismo producto. Guarda la cantidad disponible para vender y la cantidad reservada (bloqueada por pedidos en proceso). MovimientoInventario registra cada cambio en el stock como una bitácora: entradas, salidas, ajustes. InventarioSaldoHistorial es el historial temporal del stock con fechas de vigencia.")
    separador(doc)

    h3(doc, "pedido.py — Pedidos de compra")
    body(doc, "Pedido es el registro principal de una compra, con la referencia al usuario, a la dirección de entrega y los totales económicos (base imponible, IVA, total). PedidoLinea es cada artículo dentro del pedido: guarda un snapshot del precio, nombre del producto y nombre del vendedor al momento de la compra, para que el historial sea inmutable aunque el producto cambie de precio.")
    separador(doc)

    h3(doc, "pedido_vendedor.py — Subpedidos por vendedor")
    body(doc, "Un pedido que incluye productos de múltiples vendedores se divide en subpedidos, uno por vendedor. Esto permite que cada vendedor gestione sus propias líneas de forma independiente. PedidoDireccion es un snapshot de la dirección de entrega en el momento de hacer el pedido.")
    separador(doc)

    h3(doc, "pago.py — Pagos")
    body(doc, "MetodoPago lista los métodos de pago disponibles (tarjeta de crédito, transferencia, etc.). Pago registra el pago asociado a un pedido: qué método se usó, cuánto se cobró, si fue aprobado y la referencia de la transacción.")
    separador(doc)

    h3(doc, "outbox.py — Patrón Outbox transaccional")
    body(doc, "Tabla que actúa como cola de mensajes dentro de MySQL. Cuando se crea un producto o se actualiza su información, en lugar de escribir directamente en MongoDB (lo que podría fallar y dejar los datos inconsistentes), se registra un evento pendiente en esta tabla dentro de la misma transacción de MySQL. Un proceso separado (worker) lee estos eventos y los aplica en MongoDB de forma idempotente. Así se garantiza que MySQL y MongoDB siempre estén sincronizados.")
    separador(doc)

    h3(doc, "solicitud_catalogo.py — Solicitudes de catálogo")
    body(doc, "Implementa el flujo de aprobación para que los vendedores agreguen productos. Una solicitud tiene tipo (producto nuevo u oferta existente), el estado del flujo de aprobación, todos los datos propuestos por el vendedor (nombre, precio, stock, atributos) y las observaciones del administrador al revisar. También se relaciona con tablas auxiliares para las categorías propuestas y las imágenes adjuntas.")
    separador(doc)

    h2(doc, "Esquemas de Validación (app/schemas/)")
    body(doc, "Los esquemas Pydantic definen la forma exacta de los datos que entran y salen de la API. Validan automáticamente que los datos recibidos sean correctos y transforman los modelos de la base de datos en respuestas JSON.")
    separador(doc)

    h3(doc, "auth.py — Esquemas de autenticación")
    body(doc, "RegisterRequest valida los datos de registro con reglas estrictas para la contraseña (mínimo 8 caracteres, debe incluir mayúsculas, minúsculas, números y caracteres especiales). TokenResponse define el formato de la respuesta al hacer login (token de acceso y tipo). UserResponse define qué datos del usuario se devuelven en las respuestas (nunca incluye la contraseña).")
    separador(doc)

    h3(doc, "producto.py — Esquemas de productos")
    body(doc, "ProductoCreate define qué datos se necesitan para crear un producto. ProductoResponse define el formato completo de un producto al leerlo, incluyendo datos enriquecidos como el nombre del vendedor y el resumen de reseñas. ProductoListResponse envuelve una lista de productos con metadatos de paginación (página actual, total de páginas).")
    separador(doc)

    h3(doc, "checkout.py — Esquemas del proceso de compra")
    body(doc, "CheckoutRequest define los datos necesarios para completar una compra: dirección de entrega, método de pago y la lista de artículos (cada uno con el identificador de la oferta y la cantidad). CheckoutResponse confirma el resultado del pedido con su número y el total cobrado.")
    separador(doc)

    h2(doc, "Endpoints de la API (app/api/v1/)")
    separador(doc)

    h3(doc, "auth.py — Endpoints de autenticación")
    body(doc, "Cuatro endpoints: POST /auth/register crea un nuevo usuario asignándole automáticamente el rol de comprador. POST /auth/login valida las credenciales y devuelve un token JWT. GET /auth/me devuelve los datos del usuario autenticado. PUT /auth/me actualiza el perfil. La contraseña nunca se devuelve en ninguna respuesta.")
    separador(doc)

    h3(doc, "products.py — Endpoints de productos")
    body(doc, "GET /products lista el catálogo con soporte para filtros de categoría, rango de precio, búsqueda por texto y ordenamiento. GET /products/{id} devuelve el detalle completo de un producto junto con todas sus ofertas activas y el stock disponible. GET /products/images/{id} sirve las imágenes binarias almacenadas en MySQL con cabeceras de caché para evitar descargas repetidas.")
    separador(doc)

    h3(doc, "cart.py — Endpoints del carrito")
    body(doc, "El carrito del usuario se mantiene en el servidor. GET /cart devuelve el carrito enriquecido con el nombre del producto, imagen y precio actualizado. POST /cart/items agrega un artículo; si ya existe la misma oferta en el carrito, suma las cantidades en lugar de duplicar. DELETE /cart/items/{id} elimina un artículo específico.")
    separador(doc)

    h3(doc, "orders.py — Endpoints de pedidos y checkout")
    body(doc, "POST /orders/checkout es el endpoint más crítico: recibe el pedido y lo procesa en una transacción ACID completa. Una vez creado el pedido exitosamente, lanza en segundo plano la generación y envío por correo de la factura PDF para no hacer esperar al usuario. GET /orders/{id}/invoice genera y devuelve el PDF de la factura en tiempo real.")
    separador(doc)

    h3(doc, "notifications.py — Endpoints de notificaciones")
    body(doc, "Gestiona el sistema de notificaciones en tiempo real. El frontend consulta el contador de no leídas cada 60 segundos. Al abrir la bandeja de notificaciones, el frontend llama a get y luego a read-all para marcarlas como leídas.")
    separador(doc)

    h3(doc, "vendor.py — Endpoints del panel de vendedor")
    body(doc, "Todos los endpoints de este archivo requieren el rol de vendedor. GET /vendor/stats calcula los KPIs del vendedor con una consulta a MySQL. GET /vendor/orders devuelve los subpedidos del vendedor con paginación. PATCH /vendor/orders/{id}/status permite al vendedor cambiar el estado de un subpedido (por ejemplo, marcarlo como enviado).")
    separador(doc)

    h3(doc, "admin.py — Endpoints de administración")
    body(doc, "El router más extenso. Todos los endpoints requieren rol de administrador. Incluye CRUD completo de productos (con subida de imágenes y creación de ofertas iniciales), gestión de categorías y sus esquemas de atributos, gestión de usuarios y roles, visualización de todos los pedidos, reportes de ventas por vendedor y por producto, y exportación de datos a Excel. La creación de un producto nuevo es la operación más compleja: crea el documento en MongoDB, registra la referencia en MySQL, crea una oferta inicial y registra el inventario inicial, todo en una sola operación compensada.")
    separador(doc)

    h3(doc, "catalog_requests.py — Flujo de solicitudes de catálogo")
    body(doc, "Dos routers: uno para vendedores y otro para administradores. Los vendedores pueden proponer productos nuevos con imágenes y atributos, o solicitar ser oferentes de un producto existente. Los administradores pueden ver todas las solicitudes pendientes y aprobar o rechazar cada una. Al aprobar una solicitud de producto nuevo, el sistema crea automáticamente el producto en MongoDB y MySQL, la oferta y el inventario inicial. Ambas partes reciben notificaciones con el resultado.")
    separador(doc)

    h2(doc, "Servicios de Lógica de Negocio (app/services/)")
    separador(doc)

    h3(doc, "catalog_service.py — Catálogo de productos")
    body(doc, "Centraliza toda la lógica de acceso al catálogo. La función listar_productos construye la consulta a MongoDB aplicando los filtros y la paginación, luego enriquece cada resultado con las ofertas activas y el stock disponible obtenidos de MySQL. La función crear_producto inserta el documento en MongoDB y registra el evento correspondiente para mantener el historial. La función stats_catalogo ejecuta una agregación en MongoDB que calcula KPIs del catálogo en paralelo usando $facet.")
    separador(doc)

    h3(doc, "checkout_service.py — Procesamiento del pedido")
    body(doc, "Implementa la transacción completa de checkout. El proceso es: validar usuario y dirección, bloquear con SELECT FOR UPDATE las filas de ofertas e inventario (para evitar que dos usuarios compren el último producto simultáneamente), verificar stock suficiente, calcular los totales extrayendo el IVA del precio (que ya lo incluye), crear el pedido con sus líneas y subpedidos por vendedor, registrar el pago, convertir el carrito a estado 'convertido', y encolar eventos en la tabla Outbox para notificaciones.")
    separador(doc)

    h3(doc, "offer_service.py — Gestión de ofertas")
    body(doc, "Implementa el historial de precios y estados con el patrón SCD tipo 2: al cambiar el precio de una oferta, en lugar de sobreescribir el valor anterior, se cierra el período de vigencia del registro actual y se crea uno nuevo. Así se puede saber cuál era el precio en cualquier fecha pasada. También incluye funciones para seleccionar la mejor oferta de un producto (la de menor precio activa) y para validar que una oferta sea comprable antes de procesarla en checkout.")
    separador(doc)

    h3(doc, "invoice_service.py — Generación de facturas")
    body(doc, "Genera el PDF de la factura usando la librería ReportLab. Construye el documento con el encabezado de TiendaYa, los datos del cliente y su dirección, la lista de productos comprados con cantidades y precios, y el resumen de totales (base imponible, IVA y total). El PDF se genera en memoria y se devuelve directamente sin guardarlo en disco.")
    separador(doc)

    h3(doc, "outbox_service.py — Worker del patrón Outbox")
    body(doc, "Implementa el procesador de eventos que sincroniza MySQL con MongoDB. Al arrancar el servidor, lanza un hilo de fondo que cada segundo revisa si hay eventos pendientes en la tabla outbox_eventos. Si los encuentra, los toma con bloqueo (SKIP LOCKED, para que múltiples instancias del servidor no procesen el mismo evento) y los aplica en MongoDB de forma idempotente. Si un evento falla, se marca con el error y se reintenta en la siguiente iteración.")
    separador(doc)

    h3(doc, "product_history_service.py — Historial de productos")
    body(doc, "Implementa Event Sourcing para productos. Cada cambio significativo (creación, cambio de precio, cambio de descripción) se registra como un evento inmutable en MongoDB. Para reconstruir el estado de un producto en una fecha específica, se replayan todos los eventos desde el inicio hasta esa fecha, aplicando cada cambio en orden. Esto permite responder preguntas como '¿Cómo era este producto hace tres meses?'.")
    separador(doc)

    h3(doc, "variant_service.py — Variantes dinámicas")
    body(doc, "Gestiona las variantes de producto. Crea variantes en MongoDB con sus atributos específicos (por ejemplo, un laptop con 32GB de RAM y 1TB de almacenamiento) y registra su ancla de referencia en MySQL. Genera una clave única para cada combinación de atributos para detectar duplicados. Normaliza los nombres de los atributos a snake_case para consistencia.")
    separador(doc)

    h3(doc, "sku_service.py — Generación de SKUs")
    body(doc, "Genera los códigos SKU únicos para productos y ofertas. Para un producto, usa el prefijo de la categoría más un sufijo hexadecimal aleatorio, verificando en MongoDB que no exista ya. Para una oferta, combina el SKU del producto con el identificador del vendedor.")
    separador(doc)

    h3(doc, "email_service.py — Envío de correos")
    body(doc, "Envía la factura PDF por correo electrónico usando SMTP. Está diseñado para fallar de forma silenciosa: si el correo no está configurado o el servidor SMTP no responde, simplemente registra el error en el log sin interrumpir el flujo principal del sistema.")
    separador(doc)

    doc.save("C:/Users/mynit/OneDrive/Escritorio/UNIS/Segundo Semestre 2026/Bases de Datos 2/proyecto/docs/01_documentacion_codigo.docx")
    print("OK Documento 1 guardado")


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO 2 — BASES DE DATOS
# ══════════════════════════════════════════════════════════════════════════════
def crear_doc2():
    doc = Document()
    doc.core_properties.title = "Documentación de Bases de Datos - TiendaYa"

    titulo = doc.add_heading("TiendaYa — Documentación de Bases de Datos", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Diseño relacional, decisiones de arquitectura y colecciones MongoDB")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    h1(doc, "Visión General: ¿Por qué dos bases de datos?")
    body(doc, "TiendaYa usa dos motores de base de datos en paralelo, cada uno para lo que hace mejor.")
    separador(doc)
    body(doc, "MySQL (base de datos relacional) garantiza la consistencia de las transacciones comerciales. Cuando un usuario compra un producto, es fundamental que el dinero, el stock y el pedido estén perfectamente sincronizados. Si algo falla a mitad de la operación, MySQL puede revertir todos los cambios como si nada hubiera ocurrido (transacciones ACID). Esto es indispensable para cualquier sistema de comercio electrónico.")
    separador(doc)
    body(doc, "MongoDB (base de datos documental) maneja los datos del catálogo de productos. Los productos en un marketplace son inherentemente heterogéneos: un laptop tiene atributos como procesador y RAM, mientras que una camisa tiene talla y color. En MySQL, esto obligaría a crear decenas de columnas o tablas adicionales. En MongoDB, cada producto puede tener exactamente los atributos que necesita, sin restricciones de esquema fijo. Además, MongoDB facilita el almacenamiento de datos jerárquicos como el historial de eventos de un producto.")
    separador(doc)
    body(doc, "La regla de oro del diseño fue: todo lo que necesita integridad referencial y transacciones va a MySQL; todo lo que necesita flexibilidad de esquema y variación por categoría va a MongoDB.")

    doc.add_page_break()

    h1(doc, "PARTE 1 — MySQL: Esquema Relacional")
    separador(doc)

    h2(doc, "Bloque de Usuarios y Seguridad")
    separador(doc)

    h3(doc, "Tabla: roles")
    body(doc, "Almacena los tres roles del sistema: comprador, vendedor y administrador. Es una tabla de catálogo simple con solo id y nombre. Se mantiene separada de usuarios porque permite asignar múltiples roles a un mismo usuario (un administrador puede ser también comprador) y facilita agregar nuevos roles en el futuro sin tocar la tabla principal.")
    separador(doc)

    h3(doc, "Tabla: usuarios")
    body(doc, "La tabla central de identidad del sistema. Guarda el correo electrónico (que funciona como identificador único de login), el hash de la contraseña generado con bcrypt (nunca la contraseña en texto claro), nombre, apellido, teléfono, el estado de la cuenta (activo, inactivo o suspendido) y si el email fue verificado. La columna fecha_alta registra cuándo se creó la cuenta y fecha_actualizacion se actualiza automáticamente en cada cambio.")
    body(doc, "Decisión de diseño: el correo es UNIQUE NOT NULL porque es el identificador de login. La contraseña se guarda como hash de al menos 60 caracteres. El estado permite suspender una cuenta sin eliminarla, preservando el historial de compras del usuario.")
    separador(doc)

    h3(doc, "Tabla: usuario_rol")
    body(doc, "Tabla puente de la relación muchos-a-muchos entre usuarios y roles. Un usuario puede tener varios roles simultáneamente. La clave primaria compuesta (usuario_id, rol_id) evita que se asigne el mismo rol dos veces al mismo usuario.")
    separador(doc)

    h2(doc, "Bloque de Vendedores")
    separador(doc)

    h3(doc, "Tabla: vendedores")
    body(doc, "Perfil comercial extendido de los usuarios que son vendedores. Se mantiene separado de la tabla usuarios por la Tercera Forma Normal: los atributos comerciales (NIT, nombre del negocio, logo) no dependen del usuario como persona sino de su rol de vendedor. Es una relación uno-a-uno con usuarios.")
    body(doc, "El NIT tiene restricción UNIQUE porque en Guatemala es el identificador fiscal único de un negocio. El estado_verificacion permite que el administrador apruebe o rechace nuevos vendedores antes de que puedan publicar productos.")
    separador(doc)

    h2(doc, "Bloque de Direcciones")
    separador(doc)

    h3(doc, "Tabla: direcciones")
    body(doc, "Cada fila es una dirección guardada por un usuario. El tipo distingue entre dirección de envío y de facturación. La columna es_predeterminada indica cuál usar automáticamente en el checkout. La columna activa permite borrado lógico: la dirección no se elimina para no perder la referencia en pedidos históricos, solo se marca como inactiva para ocultarla al usuario.")
    body(doc, "La granularidad de pais, departamento, municipio, linea1 y linea2 responde al sistema de direcciones guatemalteco.")
    separador(doc)

    h2(doc, "Bloque de Catálogo")
    separador(doc)

    h3(doc, "Tabla: categorias")
    body(doc, "Árbol auto-referenciado de categorías. La columna categoria_padre_id apunta a otra fila de la misma tabla, lo que permite niveles ilimitados de jerarquía. Las categorías raíz (Electrónica, Ropa, etc.) tienen categoria_padre_id NULL. Las subcategorías (Computadoras, Celulares) apuntan a su categoría padre.")
    body(doc, "El slug es la versión URL de la categoría (por ejemplo, 'computadoras') y se usa como identificador en filtros y en MongoDB. El sku_prefix es el prefijo de dos a cuatro letras que se agrega a los SKUs de los productos de esa categoría, por ejemplo TECH para tecnología.")
    separador(doc)

    h3(doc, "Tabla: producto_referencias")
    body(doc, "Esta tabla fue central en la decisión de diseño híbrido. En lugar de duplicar todos los datos del producto en MySQL, solo se guarda una referencia mínima: el producto_ref (un ObjectId de MongoDB en formato hexadecimal de 24 caracteres) y la categoría principal.")
    body(doc, "Por qué este diseño: los productos tienen atributos muy variables según su categoría. Si se pusieran todos los posibles atributos en MySQL, se tendría una tabla con decenas de columnas donde la mayoría estarían vacías para cada producto. En cambio, MySQL guarda solo lo que necesita para las relaciones (la referencia), y MongoDB guarda el detalle completo. Esto sigue el principio de responsabilidad única: MySQL gestiona las transacciones, MongoDB gestiona el contenido.")
    separador(doc)

    h3(doc, "Tabla: producto_referencia_categorias")
    body(doc, "Tabla puente para que un producto pueda pertenecer a múltiples categorías. La columna es_principal indica cuál es la categoría primaria, usada para el filtrado principal del catálogo. Un laptop puede pertenecer tanto a 'Computadoras' como a 'Electrónica'.")
    separador(doc)

    h3(doc, "Tabla: producto_imagenes")
    body(doc, "Almacena las imágenes de productos directamente en MySQL como datos binarios (LONGBLOB). Esta fue una decisión deliberada para el proyecto académico: en producción real se usaría almacenamiento en la nube (S3, Cloudinary), pero para este contexto MySQL sirve como repositorio simplificado. La columna orden define en qué posición aparece cada imagen en la galería del producto.")
    separador(doc)

    h3(doc, "Tabla: producto_variante_referencias")
    body(doc, "Ancla SQL mínima para las variantes dinámicas de productos. Al igual que producto_referencias, solo guarda el identificador de la variante en MongoDB (variante_ref) y a qué producto pertenece. El detalle de la variante (qué atributos la definen) vive en MongoDB. Esta separación permite que las variantes sean completamente dinámicas sin agregar columnas a MySQL.")
    separador(doc)

    h2(doc, "Bloque de Ofertas y Precios")
    separador(doc)

    h3(doc, "Tabla: ofertas")
    body(doc, "Una oferta es la unidad comprable del marketplace: un vendedor específico vendiendo un producto específico en una variante específica a un precio determinado. Esta tabla es la que el carrito y el checkout usan como referencia, no el producto directamente.")
    body(doc, "Por qué este diseño: en un marketplace un mismo producto puede ser vendido por múltiples vendedores a diferentes precios. La oferta encapsula esa relación. El campo producto_variante_id vincula la oferta a una variante específica (o a la variante predeterminada si el producto no tiene variantes). El índice único uq_oferta_vendedor_variante garantiza que un vendedor no pueda tener dos ofertas para exactamente la misma variante del producto.")
    separador(doc)

    h3(doc, "Tabla: oferta_precios_historial")
    body(doc, "Implementa el patrón Slowly Changing Dimension tipo 2 (SCD2) para los precios. En lugar de sobreescribir el precio anterior al hacer un cambio, se cierra el período de vigencia del registro actual (se llena vigente_hasta) y se crea un nuevo registro con el precio nuevo. La columna es_vigente es un campo calculado por MySQL que vale 1 si vigente_hasta es NULL (es el precio actual) y NULL en caso contrario.")
    body(doc, "Por qué SCD2: permite responder 'cuánto costaba este producto el 15 de enero' con una simple consulta, lo que es invaluable para auditorías, disputas de precios y análisis histórico. El índice único sobre (oferta_id, es_vigente) garantiza que solo haya un precio vigente por oferta en todo momento.")
    separador(doc)

    h3(doc, "Tabla: oferta_estados_historial")
    body(doc, "El mismo patrón SCD2 pero para el estado de las ofertas (borrador, activa, pausada, descontinuada). Permite auditar cuándo una oferta fue activada, pausada o dada de baja, y por quién.")
    separador(doc)

    h2(doc, "Bloque de Inventario")
    separador(doc)

    h3(doc, "Tabla: inventario")
    body(doc, "Control de stock a nivel de oferta. El inventario se gestiona por oferta (no solo por producto) porque cada vendedor puede tener diferente stock del mismo producto en su bodega. La columna cantidad_disponible es el stock real. La columna cantidad_reservada representa unidades bloqueadas por pedidos en proceso que aún no han sido confirmados; esto evita sobrevender el mismo artículo.")
    body(doc, "El stock vendible real es cantidad_disponible menos cantidad_reservada. Durante el checkout se incrementa cantidad_reservada y en el movimiento de inventario real se decrementa cantidad_disponible.")
    separador(doc)

    h3(doc, "Tabla: movimientos_inventario")
    body(doc, "Bitácora de cada cambio en el stock: entradas (reabastecimiento), salidas (ventas), ajustes (correcciones manuales), reservas y liberaciones. Registra quién hizo el cambio y por qué. Es el historial de auditoría del inventario.")
    separador(doc)

    h3(doc, "Tabla: inventario_saldos_historial")
    body(doc, "SCD2 para los saldos de inventario. Registra cómo cambiaron las cantidades disponibles y reservadas a lo largo del tiempo.")
    separador(doc)

    h2(doc, "Bloque de Pedidos")
    separador(doc)

    h3(doc, "Tabla: pedidos")
    body(doc, "El registro principal de una compra. Contiene quién compró, en qué dirección quiere recibir, cuándo y cuánto pagó. Los campos subtotal, impuestos y total almacenan la base imponible (precio sin IVA), el IVA extraído (12%) y el precio total respectivamente. En Guatemala el IVA está incluido en el precio del producto, por lo que estos campos reflejan el desglose de un precio que ya incluía el impuesto.")
    separador(doc)

    h3(doc, "Tabla: pedido_vendedores")
    body(doc, "Cuando un pedido incluye productos de dos vendedores distintos, se crean dos subpedidos, uno por vendedor. Esto permite que cada vendedor gestione su parte de forma independiente: el vendedor A puede marcar sus artículos como enviados sin afectar al vendedor B. Cada subpedido tiene su propio estado y subtotal.")
    separador(doc)

    h3(doc, "Tabla: pedido_direcciones")
    body(doc, "Snapshot inmutable de la dirección de entrega al momento del pedido. Por qué snapshot: si el usuario cambia su dirección después de hacer el pedido, el pedido antiguo debe seguir mostrando la dirección original a la que fue enviado. Esta es la misma razón por la que pedido_lineas guarda el nombre del producto y el precio congelado.")
    separador(doc)

    h3(doc, "Tabla: pedido_lineas")
    body(doc, "Cada artículo comprado en un pedido. Guarda snapshots inmutables del producto_nombre, vendedor_nombre_snapshot, precio_unitario y sku_snapshot al momento de la compra. Si el producto cambia de nombre o precio después, el historial del pedido sigue siendo correcto. La columna subtotal_linea es precio_unitario por cantidad.")
    separador(doc)

    h2(doc, "Bloque de Pagos")
    separador(doc)

    h3(doc, "Tabla: metodos_pago")
    body(doc, "Catálogo de los métodos de pago habilitados: tarjeta de crédito, débito, transferencia, contra entrega. La columna activo permite habilitar o deshabilitar un método sin eliminarlo.")
    separador(doc)

    h3(doc, "Tabla: pagos")
    body(doc, "Registro del pago asociado a un pedido. En esta fase del proyecto es un registro simple que confirma el método, monto y estado del pago. La referencia_transaccion guarda el identificador que devolvería una pasarela de pago real.")
    separador(doc)

    h2(doc, "Bloque de Carrito")
    separador(doc)

    h3(doc, "Tabla: carritos")
    body(doc, "Cada usuario tiene un carrito activo. El estado del carrito puede ser activo (en uso), abandonado o convertido (cuando el usuario completó la compra). Tener el carrito en el servidor (y no solo en el navegador) permite recuperarlo si el usuario cierra y vuelve a abrir el navegador.")
    separador(doc)

    h3(doc, "Tabla: carrito_items")
    body(doc, "Cada artículo en el carrito referencia una oferta específica. La columna precio_al_agregar guarda el precio cuando se agregó el artículo, permitiendo mostrar al usuario si el precio cambió desde que lo agregó. El índice único uq_ci_carrito_oferta impide agregar la misma oferta dos veces; en su lugar, se actualiza la cantidad.")
    separador(doc)

    h2(doc, "Bloque de Comunicación y Eventos")
    separador(doc)

    h3(doc, "Tabla: notificaciones")
    body(doc, "Notificaciones en tiempo real para los usuarios. Cada notificación tiene un tipo, título, mensaje y si fue leída. El campo pedido_id es opcional y permite vincular la notificación con un pedido específico para mostrar un enlace directo. El índice compuesto sobre (usuario_id, leida, fecha_creacion) acelera la consulta más frecuente: las notificaciones no leídas de un usuario ordenadas por fecha.")
    separador(doc)

    h3(doc, "Tabla: outbox_eventos")
    body(doc, "Implementa el patrón Transactional Outbox para garantizar la consistencia entre MySQL y MongoDB. Cuando se crea o modifica un producto, el evento se registra en esta tabla en la misma transacción de MySQL. Un proceso worker lee los eventos pendientes y los aplica en MongoDB. Si MongoDB falla temporalmente, los eventos quedan pendientes y se procesan cuando vuelva a estar disponible. Esto garantiza que nunca se pierda información aunque uno de los sistemas falle.")
    body(doc, "Los campos intentos y ultimo_error permiten monitorear eventos que están fallando repetidamente. El campo estado evoluciona de pendiente a procesando, luego a procesado (o error si falla).")
    separador(doc)

    h2(doc, "Bloque de Solicitudes de Catálogo")
    separador(doc)

    h3(doc, "Tabla: solicitudes_catalogo")
    body(doc, "Flujo de aprobación para que los vendedores propongan productos. El campo tipo indica si es una propuesta de producto completamente nuevo o una solicitud para vender un producto que ya existe. El campo estado sigue el ciclo pendiente, aprobada o rechazada. Los campos observaciones_vendedor y observaciones_admin guardan las notas de cada parte durante el proceso de revisión. Los campos producto_ref_resultado y oferta_id_resultado guardan las referencias a los registros creados si la solicitud fue aprobada.")
    separador(doc)

    h3(doc, "Tablas: solicitud_catalogo_categorias y solicitud_catalogo_imagenes")
    body(doc, "Tablas auxiliares que almacenan las categorías propuestas e imágenes adjuntas para cada solicitud. Se mantienen separadas para cumplir con la normalización (una solicitud puede tener múltiples categorías y múltiples imágenes).")
    separador(doc)

    h2(doc, "Bloque de Reseñas")
    separador(doc)

    h3(doc, "Tabla: resenas")
    body(doc, "Las reseñas se hacen sobre el producto de referencia (no sobre la oferta de un vendedor específico), porque la calificación refleja la calidad del producto en general. El campo aprobada permite moderar las reseñas antes de publicarlas.")
    separador(doc)

    doc.add_page_break()

    h1(doc, "PARTE 2 — MongoDB: Colecciones Documentales")
    separador(doc)

    h2(doc, "Colección: productos")
    body(doc, "Es la colección principal y la razón de ser de MongoDB en este sistema. Cada documento representa un producto con toda su información: nombre, descripción, precio, estado, categorías a las que pertenece, atributos técnicos específicos de su categoría, imágenes (URLs o referencias), historial resumido de cambios y métricas de reseñas.")
    separador(doc)
    body(doc, "Por qué MongoDB para los productos: los atributos de un laptop son completamente diferentes a los de una camisa o un libro. En MySQL esto obligaría a una tabla de atributos genérica (Entity-Attribute-Value) que es muy difícil de consultar y mantener, o bien decenas de columnas vacías. MongoDB permite que cada documento tenga exactamente los campos que necesita. La validación de esquema (validationLevel: moderate) asegura que los campos fundamentales (sku, nombre, precio, categoría, estado) siempre estén presentes, pero los atributos específicos pueden variar libremente.")
    separador(doc)
    body(doc, "Campos principales de cada documento de producto:")
    bullet(doc, "sku: Código único del producto generado con el prefijo de su categoría")
    bullet(doc, "nombre y descripcion: El título y descripción que ve el comprador")
    bullet(doc, "precio: El precio actual en quetzales (ya incluye IVA)")
    bullet(doc, "estado: Si el producto es activo, borrador, inactivo o descontinuado")
    bullet(doc, "categoria: Objeto anidado con slug y nombre de la categoría principal")
    bullet(doc, "categorias: Array de todas las categorías a las que pertenece")
    bullet(doc, "atributos: Objeto libre con los atributos específicos del tipo de producto")
    bullet(doc, "disponible: Booleano calculado por el sistema según el inventario")
    bullet(doc, "imagenes: Array de referencias a las imágenes del producto")
    bullet(doc, "vendedor_id: Referencia al vendedor que creó el producto en MySQL")
    bullet(doc, "producto_ref: El mismo ObjectId del documento, usado como FK en MySQL")
    separador(doc)
    body(doc, "Índices de la colección:")
    bullet(doc, "Índice único en sku: garantiza que no haya dos productos con el mismo código")
    bullet(doc, "Índice compuesto en (categoria.slug, disponible, precio): el índice más importante. Cubre el 90% de las consultas del catálogo público que filtran por categoría, luego por disponibilidad y ordenan por precio. Sin este índice, cada visita al catálogo haría un scan completo de la colección")
    bullet(doc, "Índice de texto en nombre y descripcion: permite la búsqueda por palabras clave desde la barra de búsqueda")
    bullet(doc, "Índice en (categorias.slug, estado): para los filtros de administración que necesitan ver todos los productos de una categoría en un estado específico")
    separador(doc)

    h2(doc, "Colección: producto_eventos")
    body(doc, "Implementa el patrón Event Sourcing. Cada cambio significativo en un producto genera un evento inmutable que se agrega a esta colección. Los tipos de eventos son: PRODUCTO_CREADO, PRECIO_ACTUALIZADO, DESCRIPCION_ACTUALIZADA, DISPONIBILIDAD_CAMBIADA, ATRIBUTOS_ACTUALIZADOS y PRODUCTO_DESCONTINUADO.")
    separador(doc)
    body(doc, "Por qué Event Sourcing para el historial: el historial de cambios de un producto no puede simplemente sobreescribirse. Si se guarda solo el estado actual, se pierde la información de cómo llegó a ese estado. Con eventos inmutables, se puede reconstruir el estado exacto del producto en cualquier momento del pasado, solo aplicando los eventos en orden cronológico. Esto es útil para auditorías, análisis de precios y resolver disputas.")
    separador(doc)
    body(doc, "El índice compuesto en (producto_id, timestamp) cubre el patrón de acceso más frecuente: obtener todos los eventos de un producto ordenados por fecha para reconstruir su historial.")
    separador(doc)

    h2(doc, "Colección: categoria_esquemas")
    body(doc, "Define los atributos disponibles para cada categoría. Cada documento representa una categoría y contiene un array de atributos con nombre, etiqueta, tipo de dato (texto, número, booleano, selección), si es requerido y cuáles son las opciones válidas si es un campo de selección.")
    separador(doc)
    body(doc, "Por qué MongoDB para los esquemas: los esquemas de categorías cambian con el tiempo a medida que se agregan nuevas categorías al marketplace. Almacenarlos en MongoDB permite modificarlos sin cambiar el esquema de la base de datos relacional. El frontend lee estos esquemas para generar dinámicamente los formularios de carga de atributos. Un formulario para computadoras tendrá campos de procesador y RAM; uno para ropa tendrá talla y color.")
    separador(doc)
    body(doc, "Esta colección es la que hace posible la flexibilidad del marketplace: sin ella, agregar una nueva categoría requeriría cambios en el código del formulario. Con ella, solo hay que agregar un nuevo documento en esta colección.")
    separador(doc)

    h2(doc, "Colección: producto_variantes")
    body(doc, "Colección creada para el sistema de variantes dinámicas. Cada documento representa una variante específica de un producto, definida por una combinación única de atributos. Por ejemplo, un laptop puede tener una variante con 16GB de RAM y 512GB de almacenamiento, y otra con 32GB de RAM y 1TB de almacenamiento.")
    separador(doc)
    body(doc, "Campos principales:")
    bullet(doc, "producto_ref: El ObjectId del producto al que pertenece esta variante")
    bullet(doc, "sku_catalogo: El SKU único de esta variante en el catálogo")
    bullet(doc, "atributos: Objeto con los atributos que diferencian esta variante de las demás")
    bullet(doc, "clave_variante: Representación textual normalizada de los atributos (por ejemplo, 'almacenamiento=1TB|ram_gb=32') usada para deduplicación")
    bullet(doc, "estado: Si la variante está activa o no")
    bullet(doc, "es_predeterminada: Verdadero para la variante generada automáticamente por compatibilidad con productos sin variantes explícitas")
    separador(doc)
    body(doc, "Por qué MongoDB para las variantes: al igual que los atributos de productos, los atributos que definen una variante cambian según la categoría. En MySQL se tendría que definir por adelantado qué columnas puede tener una variante. En MongoDB, el objeto atributos puede contener cualquier combinación de campos.")
    separador(doc)
    body(doc, "Índices:")
    bullet(doc, "Índice único en (producto_ref, clave_variante): garantiza que no haya dos variantes idénticas del mismo producto")
    bullet(doc, "Índice único en sku_catalogo: garantiza unicidad del SKU de variante")

    doc.save("C:/Users/mynit/OneDrive/Escritorio/UNIS/Segundo Semestre 2026/Bases de Datos 2/proyecto/docs/02_documentacion_bases_datos.docx")
    print("OK Documento 2 guardado")


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO 3 — LLAMADAS A BASES DE DATOS
# ══════════════════════════════════════════════════════════════════════════════
def crear_doc3():
    doc = Document()
    doc.core_properties.title = "Accesos a Bases de Datos - TiendaYa"

    titulo = doc.add_heading("TiendaYa — Mapa de Accesos a las Bases de Datos", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Dónde, cómo y por qué se toca MySQL y MongoDB en cada parte del código")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    h1(doc, "Introducción: Cómo llegan las peticiones a la base de datos")
    body(doc, "Cuando el frontend hace una petición al servidor, el flujo es siempre el mismo: el endpoint de la API recibe la petición, llama a un servicio que contiene la lógica de negocio, y el servicio es quien se comunica directamente con MySQL o MongoDB. Los endpoints no tocan la base de datos directamente.")
    separador(doc)
    body(doc, "MySQL se accede a través de SQLAlchemy ORM. Cada endpoint que necesita MySQL recibe automáticamente una sesión de base de datos a través del parámetro db: Session = Depends(get_db). Esta sesión es una conexión del pool y se cierra automáticamente al terminar la petición.")
    separador(doc)
    body(doc, "MongoDB se accede a través de PyMongo. La conexión es una única instancia global (singleton) que se crea al arrancar el servidor y se mantiene viva. Los endpoints que necesitan MongoDB reciben la base de datos a través del parámetro mongo=Depends(get_mongo).")
    separador(doc)
    body(doc, "Muchas operaciones combinan ambas bases de datos en la misma petición. Por ejemplo, ver el catálogo consulta MongoDB para los datos del producto y MySQL para los precios y stock de las ofertas.")

    doc.add_page_break()

    h1(doc, "PARTE 1 — Accesos solo a MySQL")
    separador(doc)

    h2(doc, "Autenticación y perfil de usuario")
    h3(doc, "Archivo: app/api/v1/auth.py — Todos los endpoints")
    body(doc, "El endpoint POST /auth/register hace una consulta a MySQL para verificar que el correo no exista ya (SELECT sobre usuarios), y si no existe, inserta el nuevo usuario y le asigna el rol de comprador (INSERT en usuarios e INSERT en usuario_rol).")
    separador(doc)
    body(doc, "El endpoint POST /auth/login hace SELECT en usuarios filtrando por email para obtener el hash de la contraseña, lo verifica con bcrypt y si es correcto devuelve un JWT. No toca MongoDB.")
    separador(doc)
    body(doc, "El endpoint GET /auth/me hace SELECT en usuarios filtrando por el id que viene en el JWT, y también carga los roles del usuario. No toca MongoDB.")
    separador(doc)

    h2(doc, "Gestión de direcciones")
    h3(doc, "Archivo: app/api/v1/addresses.py — Todos los endpoints")
    body(doc, "Todos los endpoints de direcciones (listar, crear, actualizar, eliminar) operan exclusivamente en MySQL sobre la tabla direcciones. Cuando se crea una dirección como predeterminada, también actualiza las demás direcciones del usuario para desmarcarlas (UPDATE en direcciones). El borrado es lógico: UPDATE activa=False, no DELETE.")
    separador(doc)

    h2(doc, "Carrito de compras")
    h3(doc, "Archivo: app/api/v1/cart.py y app/services/cart_service.py")
    body(doc, "El endpoint GET /cart consulta MySQL para obtener los artículos del carrito activo del usuario. Hace JOIN entre carritos, carrito_items, ofertas, y producto_referencias para obtener el precio actualizado de cada oferta. Luego, con los producto_ref obtenidos de MySQL, consulta MongoDB para obtener el nombre e imagen de cada producto. Esta es una de las operaciones que combina ambas bases de datos.")
    separador(doc)
    body(doc, "El endpoint POST /cart/items inserta o actualiza en MySQL la tabla carrito_items. No toca MongoDB.")
    separador(doc)
    body(doc, "El endpoint DELETE /cart/items/{id} elimina la fila correspondiente de carrito_items en MySQL. No toca MongoDB.")
    separador(doc)

    h2(doc, "Pedidos")
    h3(doc, "Archivo: app/api/v1/orders.py y app/services/checkout_service.py")
    body(doc, "El endpoint GET /orders consulta MySQL haciendo SELECT en pedidos, pedido_lineas y pagos para el usuario autenticado.")
    separador(doc)
    body(doc, "El endpoint POST /orders/checkout es el más complejo. En MySQL realiza las siguientes operaciones dentro de una única transacción:")
    bullet(doc, "SELECT FOR UPDATE en ofertas para bloquear las filas y evitar condiciones de carrera")
    bullet(doc, "SELECT FOR UPDATE en inventario para bloquear el stock")
    bullet(doc, "Verifica disponibilidad: cantidad_disponible - cantidad_reservada >= cantidad solicitada")
    bullet(doc, "INSERT en pedidos con subtotal (base), impuestos (IVA extraído) y total")
    bullet(doc, "INSERT en pedido_vendedores (un registro por cada vendedor involucrado)")
    bullet(doc, "INSERT en pedido_lineas (un registro por cada artículo)")
    bullet(doc, "INSERT en pedido_direcciones (snapshot de la dirección de entrega)")
    bullet(doc, "INSERT en pagos")
    bullet(doc, "UPDATE en inventario: incrementa cantidad_reservada")
    bullet(doc, "INSERT en movimientos_inventario (registro de auditoría)")
    bullet(doc, "UPDATE en carritos: cambia estado a 'convertido'")
    bullet(doc, "INSERT en outbox_eventos: encola notificaciones para admin y vendedores")
    body(doc, "No toca MongoDB directamente durante el checkout; el Outbox worker lo hace después en segundo plano.")
    separador(doc)

    h2(doc, "Notificaciones")
    h3(doc, "Archivo: app/api/v1/notifications.py")
    body(doc, "Todos los endpoints de notificaciones (listar, contar no leídas, marcar como leídas) operan exclusivamente en MySQL sobre la tabla notificaciones. El polling del frontend cada 60 segundos hace SELECT COUNT(*) WHERE leida = 0 para obtener el número de notificaciones pendientes, una consulta muy liviana gracias al índice compuesto en (usuario_id, leida, fecha_creacion).")
    separador(doc)

    h2(doc, "Panel de vendedor — estadísticas y pedidos")
    h3(doc, "Archivo: app/api/v1/vendor.py")
    body(doc, "El endpoint GET /vendor/stats hace varias consultas a MySQL: COUNT de subpedidos del vendedor, SUM de ingresos filtrando por las ofertas del vendedor en pedido_lineas, y COUNT de subpedidos en estado 'pendiente' o 'confirmado'. Todo en MySQL, sin tocar MongoDB.")
    separador(doc)
    body(doc, "El endpoint GET /vendor/orders consulta MySQL con JOIN entre pedido_vendedores, pedidos, pedido_lineas y usuarios para obtener los subpedidos del vendedor con sus líneas de productos y datos del comprador. No toca MongoDB.")
    separador(doc)

    doc.add_page_break()

    h1(doc, "PARTE 2 — Accesos solo a MongoDB")
    separador(doc)

    h2(doc, "Esquemas de categorías")
    h3(doc, "Archivo: app/api/v1/categories.py")
    body(doc, "El endpoint GET /categories/{slug}/schema consulta exclusivamente MongoDB en la colección categoria_esquemas buscando por categoria_slug. Devuelve la lista de atributos que el frontend usa para generar el formulario de atributos del producto. No toca MySQL.")
    separador(doc)

    h2(doc, "Historial de productos")
    h3(doc, "Archivo: app/services/product_history_service.py")
    body(doc, "Todas las funciones de historial de productos operan exclusivamente en MongoDB. La función registrar_evento inserta un nuevo documento en producto_eventos. La función reconstruir_estado hace una query con filtro por producto_id y rango de fechas en producto_eventos, luego aplica cada evento en orden para reconstruir el estado. La función obtener_historial hace SELECT de todos los eventos de un producto ordenados por timestamp.")
    separador(doc)

    doc.add_page_break()

    h1(doc, "PARTE 3 — Operaciones que combinan MySQL y MongoDB")
    body(doc, "Estas son las operaciones más importantes del sistema porque requieren mantener la consistencia entre ambas bases de datos.")
    separador(doc)

    h2(doc, "Catálogo de productos — Lectura")
    h3(doc, "Archivo: app/services/catalog_service.py — función listar_productos")
    body(doc, "Esta función combina ambas bases de datos en cada llamada al catálogo público. El proceso es:")
    bullet(doc, "Paso 1 (MongoDB): Consulta la colección productos aplicando los filtros de categoría, texto de búsqueda, rango de precio y disponibilidad. Obtiene los documentos paginados con toda la información del producto")
    bullet(doc, "Paso 2 (MySQL): Con los producto_ref obtenidos de MongoDB, hace una query a MySQL buscando las ofertas activas para esos productos (JOIN entre producto_referencias y ofertas) y el stock disponible (JOIN con inventario)")
    bullet(doc, "Paso 3 (enriquecimiento): Combina los resultados en Python: cada documento de MongoDB se enriquece con el precio de la oferta más económica y la disponibilidad de inventario obtenidos de MySQL")
    body(doc, "Por qué esta estrategia: MongoDB tiene los datos ricos del producto (descripción, atributos) y los índices optimizados para filtrado del catálogo. MySQL tiene el precio actualizado y el stock en tiempo real con bloqueos transaccionales. Separar las responsabilidades permite que cada base de datos haga lo que mejor sabe.")
    separador(doc)

    h2(doc, "Detalle de producto — Lectura")
    h3(doc, "Archivo: app/services/catalog_service.py — función obtener_producto")
    body(doc, "Similar al catálogo pero para un solo producto. Consulta MongoDB para el documento completo del producto, luego MySQL para todas las ofertas activas (con precio, vendedor, estado) y el inventario de cada oferta. También carga las reseñas del producto desde MySQL (tabla resenas) para calcular el promedio y la distribución de calificaciones. El resultado final combina datos de las tres fuentes.")
    separador(doc)

    h2(doc, "Creación de producto — Escritura")
    h3(doc, "Archivo: app/api/v1/admin.py — endpoint POST /admin/products")
    body(doc, "La creación de un producto es la operación más compleja del sistema porque debe mantener consistencia entre MySQL y MongoDB. El proceso en orden:")
    bullet(doc, "Paso 1 (MySQL): Abre una transacción. Crea el registro en producto_referencias con la categoría")
    bullet(doc, "Paso 2 (MongoDB): Inserta el documento del producto con todos sus atributos, usando el id de MySQL como producto_ref")
    bullet(doc, "Paso 3 (MongoDB): Registra el evento PRODUCTO_CREADO en producto_eventos")
    bullet(doc, "Paso 4 (MySQL, dentro de la transacción): Crea la oferta inicial en ofertas, el inventario inicial en inventario, el historial de precio en oferta_precios_historial y el historial de estado en oferta_estados_historial")
    bullet(doc, "Paso 5 (MySQL, dentro de la transacción): Inserta el evento en outbox_eventos para sincronización posterior")
    bullet(doc, "Paso 6: Confirma la transacción MySQL (commit)")
    body(doc, "Si algo falla en los pasos de MongoDB (paso 2 o 3), se intenta revertir. Si algo falla en MySQL (paso 4 en adelante), la transacción se revierte automáticamente y se elimina el documento de MongoDB para compensar. Este patrón se llama compensación y es necesario porque MongoDB y MySQL no pueden participar en la misma transacción ACID.")
    separador(doc)

    h2(doc, "Variantes dinámicas — Creación")
    h3(doc, "Archivo: app/services/variant_service.py — función create_variant")
    body(doc, "La creación de una variante también combina ambas bases de datos. Primero inserta el documento de la variante en la colección producto_variantes de MongoDB, obteniendo su ObjectId. Luego, en MySQL, inserta una fila en producto_variante_referencias con ese ObjectId como variante_ref. La clave de deduplicación (clave_variante) garantiza que no se creen dos variantes idénticas en MongoDB gracias al índice único.")
    separador(doc)

    h2(doc, "Solicitudes de catálogo — Aprobación")
    h3(doc, "Archivo: app/api/v1/catalog_requests.py — endpoint PATCH admin/catalog-requests/{id}/review")
    body(doc, "Cuando un administrador aprueba una solicitud de producto nuevo, el sistema ejecuta:")
    bullet(doc, "MySQL: Crea producto_referencias, ofertas, inventario, oferta_precios_historial")
    bullet(doc, "MongoDB: Inserta el documento de producto en productos y el evento PRODUCTO_CREADO en producto_eventos")
    bullet(doc, "MySQL: Actualiza solicitudes_catalogo con estado='aprobada', referencias a los registros creados y la fecha de revisión")
    bullet(doc, "MySQL: Inserta notificación para el vendedor en la tabla notificaciones")
    body(doc, "Si el administrador rechaza, solo se actualiza la solicitud en MySQL y se crea la notificación. No se toca MongoDB.")
    separador(doc)

    h2(doc, "Outbox Worker — Sincronización asíncrona")
    h3(doc, "Archivo: app/services/outbox_service.py — función start_outbox_worker")
    body(doc, "El worker es un hilo de fondo que corre permanentemente mientras el servidor está activo. Cada segundo revisa MySQL buscando eventos en estado 'pendiente' en la tabla outbox_eventos. Para cada evento encontrado:")
    bullet(doc, "MySQL: Actualiza el estado del evento a 'procesando' usando SELECT FOR UPDATE SKIP LOCKED (esto evita que dos instancias del servidor procesen el mismo evento)")
    bullet(doc, "Según el tipo de evento: aplica la operación correspondiente en MongoDB. Por ejemplo, para un evento de tipo PRODUCTO_ACTUALIZADO hace $set en la colección productos")
    bullet(doc, "MySQL: Si fue exitoso, actualiza el estado a 'procesado' y registra la fecha")
    bullet(doc, "MySQL: Si falló, actualiza el estado a 'error', incrementa el contador de intentos y registra el mensaje de error")
    body(doc, "Este patrón garantiza consistencia eventual: aunque MongoDB tarde en actualizarse, eventualmente siempre recibirá los cambios registrados en MySQL. El SKIP LOCKED es clave para que el sistema escale a múltiples instancias del servidor sin conflictos.")
    separador(doc)

    h2(doc, "Migración de variantes — Script apply_dynamic_variants.py")
    h3(doc, "Archivo: backend/scripts/apply_dynamic_variants.py")
    body(doc, "Este script de migración es el ejemplo más completo de uso combinado de ambas bases de datos. Para migrar los productos existentes al nuevo sistema de variantes dinámicas:")
    bullet(doc, "MySQL: Lee todos los registros de producto_referencias")
    bullet(doc, "MongoDB: Para cada producto_ref, busca el documento en la colección productos para obtener el SKU")
    bullet(doc, "MongoDB: Crea un documento de variante predeterminada en producto_variantes con clave_variante='__default__' (la variante que representa 'sin variante específica')")
    bullet(doc, "MySQL: Inserta el registro correspondiente en producto_variante_referencias con el ObjectId de la variante recién creada")
    bullet(doc, "MySQL: Actualiza todas las ofertas existentes para vincularlas con su variante predeterminada (UPDATE en ofertas SET producto_variante_id = pvr.id)")
    bullet(doc, "MongoDB: Crea los índices únicos en producto_variantes")
    body(doc, "El script es idempotente: si se ejecuta dos veces, no duplica datos gracias al uso de INSERT IGNORE en MySQL y find_one_and_update con upsert en MongoDB.")
    separador(doc)

    doc.add_page_break()

    h1(doc, "PARTE 4 — Resumen de archivos y su relación con las bases de datos")
    separador(doc)

    h2(doc, "Archivos que usan SOLO MySQL")
    bullet(doc, "app/api/v1/auth.py — usuarios, usuario_rol")
    bullet(doc, "app/api/v1/addresses.py — direcciones")
    bullet(doc, "app/api/v1/notifications.py — notificaciones")
    bullet(doc, "app/api/v1/vendor.py — pedido_vendedores, pedido_lineas, ofertas")
    bullet(doc, "app/services/offer_service.py — ofertas, oferta_precios_historial, oferta_estados_historial")
    bullet(doc, "app/services/offer_history_service.py — oferta_estados_historial, inventario_saldos_historial")
    bullet(doc, "app/services/invoice_service.py — lee de pedidos ya guardados en MySQL")
    bullet(doc, "app/services/email_service.py — no accede a BD directamente")
    bullet(doc, "app/services/image_service.py — no accede a BD directamente")
    separador(doc)

    h2(doc, "Archivos que usan SOLO MongoDB")
    bullet(doc, "app/services/product_history_service.py — producto_eventos")
    bullet(doc, "app/api/v1/categories.py (esquemas) — categoria_esquemas")
    bullet(doc, "database/mongo/01_init_collections.js — crea colecciones con validación de esquema")
    bullet(doc, "database/mongo/02_indexes.js — crea índices en todas las colecciones")
    separador(doc)

    h2(doc, "Archivos que combinan MySQL y MongoDB")
    bullet(doc, "app/services/catalog_service.py — LECTURA: MongoDB para datos del producto, MySQL para precios y stock")
    bullet(doc, "app/api/v1/admin.py — ESCRITURA: MySQL para referencias y transacciones, MongoDB para documentos")
    bullet(doc, "app/api/v1/catalog_requests.py — APROBACIÓN: crea en ambas bases simultáneamente con compensación")
    bullet(doc, "app/api/v1/cart.py — LECTURA: MySQL para el carrito, MongoDB para nombres e imágenes")
    bullet(doc, "app/services/checkout_service.py — ESCRITURA: solo MySQL directamente; MongoDB se actualiza por Outbox")
    bullet(doc, "app/services/outbox_service.py — SINCRONIZACIÓN: lee MySQL (outbox_eventos), escribe en MongoDB")
    bullet(doc, "app/services/variant_service.py — ESCRITURA: MongoDB para la variante, MySQL para la referencia")
    bullet(doc, "app/core/db_mongo.py — CONFIGURACIÓN: crea índices en MongoDB al arrancar")
    bullet(doc, "backend/scripts/apply_dynamic_variants.py — MIGRACIÓN: lee de ambas, escribe en ambas")
    separador(doc)

    h2(doc, "El punto de conexión entre MySQL y MongoDB: el campo producto_ref")
    body(doc, "El campo producto_ref es el puente entre las dos bases de datos. Es el ObjectId de MongoDB (una cadena hexadecimal de 24 caracteres como '6a8dafce3afc5183e9e7ac10') que se almacena como CHAR(24) en la columna producto_ref de las tablas MySQL producto_referencias y ofertas.")
    separador(doc)
    body(doc, "Cuando el backend necesita datos de un producto, primero consulta MySQL para obtener el producto_ref (por ejemplo, al listar los artículos del carrito), y luego usa ese producto_ref para buscar el documento en MongoDB. En sentido inverso, cuando se crea un producto en MongoDB, su ObjectId se guarda inmediatamente en MySQL para que otras tablas puedan hacer referencia a él.")
    separador(doc)
    body(doc, "De la misma forma, variante_ref es el puente para las variantes: el ObjectId de la variante en MongoDB se almacena en la columna variante_ref de la tabla producto_variante_referencias en MySQL.")

    doc.save("C:/Users/mynit/OneDrive/Escritorio/UNIS/Segundo Semestre 2026/Bases de Datos 2/proyecto/docs/03_accesos_bases_datos.docx")
    print("OK Documento 3 guardado")


if __name__ == "__main__":
    import sys, io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    print("Generando documentos...")
    crear_doc1()
    crear_doc2()
    crear_doc3()
    print("\nListo. Los tres documentos están en: docs/")
